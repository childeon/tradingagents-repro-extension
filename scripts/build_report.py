"""
Project report — 5 pages prose + page 6 references + appendix.
12pt TNR, 1.15 line spacing, 4pt paragraph gap, 1" margins.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
OUT  = os.path.join(BASE, "Reproducing and Evaluating TradingAgents_v2.docx")

doc = Document()

sec = doc.sections[0]
sec.page_width  = Inches(8.5);  sec.page_height  = Inches(11)
sec.top_margin  = sec.bottom_margin = Inches(1.0)
sec.left_margin = sec.right_margin  = Inches(1.0)

ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(12)
ns.paragraph_format.space_before = Pt(0)
ns.paragraph_format.space_after  = Pt(0)
ns.paragraph_format.line_spacing = 1.15


# ── helpers ───────────────────────────────────────────────────────────────────

def sp(para, before=0, after=4, line=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing = line


def body(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=4, size=12,
         bold=False, italic=False):
    p = doc.add_paragraph()
    sp(p, before=before, after=after)
    p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    return p


def heading(text, level=1):
    p = doc.add_paragraph()
    sp(p, before=8 if level == 1 else 5, after=2, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.font.bold = True; r.font.underline = (level == 1)
    return p


def page_break():
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line=1.0)
    run = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    run._r.append(br)


def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color); tcPr.append(shd)


def tcell(c, text, bold=False, center=False, white=False, size=11):
    para = c.paragraphs[0]
    sp(para, before=1, after=1, line=1.0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    run.font.bold = bold
    if white: run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def make_table(col_headers, rows, hi=0):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(col_headers))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(col_headers):
        tcell(tbl.rows[0].cells[ci], h, bold=True, center=(ci>0), white=True)
    shade_row(tbl.rows[0], '1E3A5F')
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            tcell(tbl.rows[ri+1].cells[ci], val, bold=(ri==hi and ci==0), center=(ci>0))
        if ri == hi:
            shade_row(tbl.rows[ri+1], 'D9E8F5')
    return tbl


def caption(text):
    p = doc.add_paragraph()
    sp(p, before=3, after=8, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.italic = True


# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
sp(p, before=0, after=2, line=1.0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Reproducing and Evaluating TradingAgents: A Multi-Agent LLM Trading Framework")
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.font.bold = True

body("Yujia Zhang, Ruiyan Li, Yeyuxi Yi",
     align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)

# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading("1. Introduction")

body(
    "The application of large language models to financial trading has attracted "
    "significant research interest, yet most existing work treats the LLM as a monolithic "
    "predictor rather than a reasoning agent embedded in a larger system. FinGPT "
    "(Yang et al., 2023) fine-tunes open-source LLMs on financial corpora for sentiment "
    "classification and return forecasting; FinMem (Yu et al., 2023) adds layered memory "
    "to a single-agent trader; FinAgent (Zhang et al., 2024) augments the pipeline with "
    "multimodal tool use. Deep reinforcement learning approaches (Pricope, 2021) optimize "
    "policy networks directly on price sequences but cannot incorporate unstructured "
    "textual signals. None of these approaches model deliberation among agents or "
    "assign explicit risk-management responsibilities to a dedicated component."
)
body(
    "TradingAgents: Multi-Agents LLM Financial Trading Framework (Tauric Research, "
    "arXiv 2412.20138, December 2024) addresses this gap by assembling a team of "
    "specialized LLM agents that collaborate, debate, and filter each other's outputs "
    "before generating trading orders. The system attracted over 53,000 GitHub stars "
    "within weeks of release, motivating an independent reproducibility study. We chose "
    "this paper because (i) its modular design admits controlled ablation experiments, "
    "(ii) financial performance metrics provide unambiguous evaluation criteria, and "
    "(iii) the adversarial debate mechanism is a novel LLM design pattern whose practical "
    "value under real market conditions deserves empirical scrutiny."
)

# ════════════════════════════════════════════════════════════════════════════
# 2. FRAMEWORK OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
heading("2. Framework Overview")
heading("2.1 Task", level=2)

body(
    "Given a target ticker and date range, the framework produces a daily buy, sell, or "
    "hold decision with a proposed position size, relying entirely on LLM agents to "
    "retrieve, interpret, and synthesize raw inputs—no hand-coded signals are used at "
    "inference time. Performance is measured by cumulative return, Sharpe ratio, and "
    "maximum drawdown over the backtest window."
)

heading("2.2 Architecture", level=2)

body(
    "The pipeline has four sequential stages. At the Analyst stage, four specialized "
    "agents process distinct information channels. The Market Analyst computes technical "
    "indicators (MACD, RSI, Bollinger Bands, moving averages) from Yahoo Finance price "
    "data. The Social Media Analyst classifies retail sentiment from Reddit and Twitter "
    "posts. The News Analyst extracts event-driven signals—earnings surprises, rating "
    "changes, regulatory news—from Bloomberg, Yahoo, EODHD, FinnHub, and Reddit. The Fundamentals "
    "Analyst processes SEC filings, financial statements, and insider transaction records "
    "to assess balance-sheet health and insider conviction. Each agent produces a "
    "domain-specific natural language summary via a separate LLM call with a "
    "role-specific system prompt."
)
body(
    "The four summaries pass to the Researcher stage, which consists of a Bullish and a "
    "Bearish agent. Both agents read all analyst reports and construct opposing investment "
    "arguments; they then engage in structured multi-round debate, each round consisting "
    "of a rebuttal and a revised position. This adversarial structure is motivated by "
    "Du et al. (2023), who demonstrate that multi-agent debate substantially reduces "
    "LLM hallucinations and improves factual accuracy: when models must defend positions "
    "against active criticism, they are less likely to anchor prematurely on the first "
    "plausible interpretation. A Research Manager agent synthesizes the debate transcript "
    "into a consolidated thesis. The thesis passes to the Trader agent, which determines "
    "order timing and sizing using OpenAI o1's extended chain-of-thought reasoning. "
    "Finally, a Risk Management layer—comprising Aggressive, Neutral, and Conservative "
    "sub-agents—evaluates the proposed order against current portfolio exposure and may "
    "reduce position size, decline to execute, or override the trade direction before a "
    "Fund Manager produces the final allocation. Each trading day requires approximately "
    "eleven LLM calls and over twenty tool or data-retrieval calls."
)

heading("2.3 Contribution and Related Work", level=2)

body(
    "The paper's primary contribution is the combination of task decomposition, "
    "adversarial debate, and explicit risk control within a single framework. Prior "
    "multi-agent LLM systems—AutoGen (Wu et al., 2023), MetaGPT (Hong et al., 2023), "
    "ChatDev (Qian et al., 2023)—demonstrate the value of assigning specialized roles "
    "to communicating agents for software engineering and complex reasoning, but evaluate "
    "on qualitative task completion rather than quantitative financial performance. "
    "TradingAgents applies this paradigm to a domain with objective metrics and real "
    "economic stakes. The bull/bear debate is a structural advance over ensemble methods, "
    "which average outputs without requiring models to confront and refute each other's "
    "reasoning. The three-tier risk layer reflects the empirical finding that position "
    "sizing and downside protection determine realized long-run returns independently of "
    "raw signal quality (Pricope, 2021)."
)

# ════════════════════════════════════════════════════════════════════════════
# 3. REPRODUCTION METHODOLOGY
# ════════════════════════════════════════════════════════════════════════════
heading("3. Reproduction Methodology")
heading("3.1 Scope and Model Substitution", level=2)

body(
    "The original paper evaluates three tickers over three months. At eleven LLM calls "
    "per trading day, this implies roughly 1,500 calls per ticker-month. Our initial "
    "test on two tickers over one month, using the original GPT-4o-mini and o1-mini "
    "configuration, required over three hours of wall-clock time. We reduced scope to "
    "AMZN over January 2024 (21 trading days) and substituted Gemini 2.5 Flash Lite "
    "for lower latency and cost. AMZN was chosen for its high liquidity and abundant "
    "multi-source data coverage. We acknowledge that the model substitution introduces "
    "a cross-study confound; within-study comparisons across ablation conditions remain "
    "valid because the model is held constant across all variants."
)

heading("3.2 Engineering Challenges", level=2)

body(
    "Challenge 1 — Stale fundamentals data. The Fundamentals Analyst fetches the most "
    "recent SEC filing at run time via the EDGAR API. When executed in 2026, the API "
    "returned 2025–26 10-Q filings, introducing forward-looking bias into a 2024 "
    "backtest. The root cause is that the original code does not parameterize the "
    "filing query by a user-specified as-of date; it was written and run in late 2024 "
    "when 'most recent' was temporally consistent with the backtest window. We manually "
    "downloaded AMZN's Q3 2023 10-Q from EDGAR and cached it as a static local file. "
    "Since all 21 backtest days fall within the same fiscal quarter, this single filing "
    "is the correct fundamental input for every day, and caching also eliminates 21 "
    "redundant API calls per run."
)
body(
    "Challenge 2 — News API ordering bug and data pipeline gap. The paper's own "
    "experiments used a curated multi-source dataset—Bloomberg, Yahoo, EODHD, FinnHub, "
    "and Reddit—that was not released with the open-source code. The released code "
    "instead defaults to the free Yahoo Finance (yfinance) API as a simplified "
    "alternative. This substitution introduces a data quality gap between the paper's "
    "results and any open-source replication. Compounding this, the yfinance news "
    "endpoint returns articles in reverse-chronological order and terminates after a "
    "fixed record count; when called in 2026 with a January 2024 target window, the "
    "response contains only 2025–26 articles, all of which are discarded by the date "
    "filter, leaving the agent with empty input. We replaced the live feed with a "
    "pre-filtered static dataset from the GDELT Project, extracting AMZN-relevant "
    "articles from December 2023 through January 2024 and formatting them to match the "
    "schema expected by the News Analyst agent. The Social Media Analyst's dependency "
    "on a locally cached Reddit dataset—not released with the codebase—was replaced "
    "with a Kaggle dataset of r/wallstreetbets and r/stocks submissions filtered to "
    "AMZN mentions over the same period."
)
body(
    "Challenge 3 — Latency. The sequential chain of LLM and retrieval calls produced "
    "non-trivial per-day execution times even with Flash Lite. Beyond fundamentals "
    "caching, we pre-processed the GDELT and Reddit datasets into condensed, agent-ready "
    "summaries offline, reducing input token counts and eliminating context-overflow risk "
    "on high-news-volume days. A key limitation of these substitutions is that the "
    "resulting system is non-production-ready: it cannot generalize to new time periods "
    "without manual data preparation. All results should be interpreted as a "
    "reproducibility study under controlled, historically frozen data conditions."
)

# ════════════════════════════════════════════════════════════════════════════
# 4. EXPERIMENTAL RESULTS
# ════════════════════════════════════════════════════════════════════════════
heading("4. Experimental Results")
heading("4.1 Replication", level=2)

body(
    "We compare the full TradingAgents architecture against five baseline strategies on "
    "AMZN over January 2024: Buy-and-Hold, MACD, KDJ+RSI, Zero Mean Reversion (ZMR), "
    "and SMA (10/30). Full metrics are reported in Appendix Table A1. TradingAgents "
    "achieves a cumulative return of 4.41% (annualized: 67.84%) with a Sharpe ratio "
    "of 3.90 and a maximum drawdown of −3.57%, outperforming all baselines on both "
    "absolute and risk-adjusted return. Buy-and-Hold is the strongest competitor at "
    "3.51% return and Sharpe 2.40, reflecting the genuinely bullish environment for "
    "AMZN in January 2024. KDJ+RSI achieves a comparable Sharpe ratio (3.54) but lower "
    "cumulative return (3.26%), indicating similar risk-adjusted efficiency but less "
    "absolute upside capture. MACD lags at 0.96% return and Sharpe 0.82; ZMR and "
    "SMA (10/30) do not generate signals and return zero throughout the month."
)
body(
    "The equity curve provides additional texture. TradingAgents exhibits consistent, "
    "smooth growth with minimal drawdown during the mid-month consolidation (approximately "
    "January 8–15), consistent with the risk management layer moderating exposure when "
    "analyst signals are mixed. KDJ+RSI shows more volatile intramonth behavior with a "
    "sharper late-month recovery, reflecting the more reactive nature of rule-based "
    "oscillators. These patterns suggest that the multi-agent framework's value lies not "
    "only in end-of-period return but in producing more stable intramonth trajectories—a "
    "property that would compound favorably over longer evaluation horizons. Our "
    "replication is qualitatively consistent with the paper's claims: TradingAgents "
    "outperforms technical baselines. Direct quantitative comparison with the original "
    "paper's figures is precluded by our model substitution and modified data pipeline."
)

heading("4.2 Extension: Ablation Study", level=2)

body(
    "We evaluate two ablation variants on the same AMZN January 2024 window. In the "
    "No Research Debate variant, the bull/bear debate is removed: analyst reports are "
    "passed directly to a Research Manager that consolidates them without adversarial "
    "exchange. In the No Risk Layer variant, the Risk Management and Fund Manager layers "
    "are removed entirely, executing the Trader's output as the final order. Full results "
    "are in Appendix Table A2."
)
body(
    "Removing the debate layer unexpectedly improves performance: the No Research Debate "
    "variant achieves 6.31% cumulative return and Sharpe 5.73, compared to 4.41% and "
    "3.90 for the full architecture. January 2024 was a period of sustained upward "
    "momentum for AMZN; analyst signals were consistently bullish and directionally "
    "correct. The Bearish agent, however, is structurally required to argue against the "
    "dominant signal regardless of its strength—it cannot abstain or concede. In a "
    "strongly trending market, this produces a weakly supported counterargument that "
    "the Research Manager partially incorporates, causing the Trader to hedge "
    "unnecessarily and leave upside on the table. Without the debate, the Research "
    "Manager receives uncontested bullish inputs and acts on them more fully. This "
    "observation aligns with a known limitation identified by Du et al. (2023): a "
    "single round of debate may produce noise rather than convergence if the number of "
    "exchange rounds is insufficient for genuine reconciliation of views. The debate "
    "mechanism is likely to add value in range-bound or high-uncertainty markets where "
    "true signal disagreement is informative, but appears counterproductive when "
    "directional momentum is strong."
)
body(
    "Removing the risk layer produces the opposite effect: cumulative return falls to "
    "−1.33% (Sharpe −0.77, max drawdown −4.81%), the worst result across all "
    "configurations. Without position-level filtering, the Trader's overconfident bets "
    "are executed at full size, and small losses compound. The risk layer's conservative "
    "sub-agent plays a particularly important role: by clipping extreme position sizes "
    "during pullbacks, it prevents individual losing trades from meaningfully denting "
    "cumulative return. This result empirically confirms the foundational trading "
    "principle that signal generation and risk control are complementary rather than "
    "substitutable components of a trading system (Pricope, 2021); strong directional "
    "signals without position sizing discipline reliably produce worse realized outcomes."
)

# ════════════════════════════════════════════════════════════════════════════
# 5. DISCUSSION AND CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
heading("5. Discussion and Conclusion")

body(
    "The ablation results reveal a sharp asymmetry: the risk management layer is "
    "load-bearing and its removal is catastrophic, while the debate layer's contribution "
    "is regime-dependent and its removal is beneficial under strongly trending conditions. "
    "This suggests that the full TradingAgents architecture is not universally optimal. "
    "An adaptive design that gates the debate mechanism on a measure of inter-agent "
    "signal disagreement—activating it when analyst outputs are genuinely conflicted and "
    "suppressing it when signals are strongly aligned—could outperform both the full "
    "architecture and the no-debate ablation across diverse market regimes. The "
    "single-round debate limitation noted by Du et al. (2023) further suggests that "
    "extending to iterative multi-round exchange would allow the bear's arguments to be "
    "properly refuted rather than averaged into the bull's position."
)
body(
    "Several limitations temper these conclusions. The evaluation window is short—21 "
    "days in a positive-momentum environment for AMZN—making it difficult to assess "
    "how the architecture would perform in bear markets or high-volatility regimes. "
    "The model substitution (Gemini Flash Lite for GPT-4o-mini/o1-mini) and data "
    "substitutions (GDELT/Kaggle for live APIs) prevent direct numerical comparison "
    "with the paper's results. The metric set does not capture trading cost sensitivity, "
    "slippage, or tail risk under stress. More broadly, the reproducibility challenges "
    "we encountered—stale API responses, date-filtering bugs that surface only when "
    "running historical backtests long after development, and inaccessible locally cached "
    "datasets—are informative for the agent-systems community. Systems that implicitly "
    "assume 'most recent API response ≈ target backtest date' become silently invalid "
    "as time passes. Making data retrieval explicitly date-parameterized is a "
    "straightforward fix that would substantially improve the long-run reproducibility "
    "of this class of LLM-based agent systems."
)
body(
    "Despite these limitations, our study supports TradingAgents' central claim: a "
    "structured multi-agent pipeline with explicit risk control outperforms standard "
    "technical trading strategies on both absolute return and Sharpe ratio. The "
    "framework demonstrates that distributing cognition across specialized agents with "
    "defined interaction protocols—rather than asking a single model to jointly handle "
    "data retrieval, signal synthesis, adversarial reasoning, and portfolio management"
    "—can produce measurably better trading outcomes. Risk management is the "
    "indispensable component; structured debate is conditionally valuable. A more "
    "streamlined design that pairs strong risk control with selective, multi-round "
    "debate would be a productive direction for future work."
)

# ════════════════════════════════════════════════════════════════════════════
# PAGE 6 — REFERENCES
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading("References")

refs = [
    ("Du, Y., Li, S., Torralba, A., Tenenbaum, J. B., & Mordatch, I. (2023). "
     "Improving factuality and reasoning in language models through multiagent debate. "
     "arXiv:2305.14325."),
    ("Hong, S., Zhuge, M., Chen, J., et al. (2023). MetaGPT: Meta programming for a "
     "multi-agent collaborative framework. arXiv:2308.00352."),
    ("Pricope, T. V. (2021). Deep reinforcement learning in quantitative algorithmic "
     "trading: A review. arXiv:2106.00123."),
    ("Qian, C., Liu, W., Liu, H., et al. (2023). ChatDev: Communicative agents for "
     "software development. arXiv:2307.07924."),
    ("Tauric Research. (2024). TradingAgents: Multi-agents LLM financial trading "
     "framework. arXiv:2412.20138."),
    ("Wu, Q., Bansal, G., Zhang, J., et al. (2023). AutoGen: Enabling next-gen LLM "
     "applications via multi-agent conversation. arXiv:2308.08155."),
    ("Yang, H., Liu, X.-Y., & Wang, C. D. (2023). FinGPT: Open-source financial large "
     "language models. arXiv:2306.06031."),
    ("Yu, L., et al. (2023). FinMem: A performance-enhanced LLM trading agent with "
     "layered memory and character design. arXiv:2311.13743."),
    ("Zhang, Z., et al. (2024). A multimodal foundation agent for financial trading: "
     "Tool-augmented, diversified, and generalist. arXiv:2402.18485."),
]

for ref in refs:
    p = doc.add_paragraph()
    sp(p, before=0, after=4, line=1.0)
    p.paragraph_format.left_indent       = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ════════════════════════════════════════════════════════════════════════════
page_break()

p = doc.add_paragraph()
sp(p, before=0, after=4, line=1.0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Appendix")
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.font.bold = True

heading("Table A1: Replication — TradingAgents vs. Baselines (AMZN, January 2024)", level=2)
make_table(
    ["Strategy", "Cum. Return (%)", "Ann. Return (%)", "Sharpe Ratio", "Max DD (%)"],
    [
        ("TradingAgents",  "4.41",  "67.84",  "3.90",  "−3.57"),
        ("Buy & Hold",     "3.51",  "51.37",  "2.40",  "−3.76"),
        ("KDJ + RSI",      "3.26",  "46.92",  "3.54",  "−1.34"),
        ("MACD",           "0.96",  "12.10",  "0.82",  "−3.76"),
        ("ZMR",            "0.00",  "0.00",   "0.00",   "0.00"),
        ("SMA (10/30)",    "0.00",  "0.00",   "0.00",   "0.00"),
    ], hi=0
)
caption("Table A1. All strategies evaluated on AMZN, January 2–31, 2024 (21 trading days). "
        "Highlighted = TradingAgents full architecture.")

heading("Table A2: Ablation Study (AMZN, January 2024)", level=2)
make_table(
    ["Configuration", "Cum. Return (%)", "Sharpe Ratio", "Max DD (%)", "vs. Full"],
    [
        ("Full Architecture",      "4.41",  "3.90",  "−3.57",  "—"),
        ("No Research Debate",     "6.31",  "5.73",  "−3.76",  "↑ Better"),
        ("No Risk Layer",          "−1.33", "−0.77", "−4.81",  "↓ Much worse"),
        ("Buy & Hold (reference)", "3.51",  "2.40",  "−3.76",  "—"),
    ], hi=0
)
caption("Table A2. Ablation results. Buy & Hold included as passive reference.")

doc.save(OUT)
print("Saved:", OUT)
